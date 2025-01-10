from PyPDF2 import PdfReader
from docx import Document
import io
from PIL import Image as PILImage
import streamlit as st


@st.cache_resource
def get_docx_text(docx_file):
    doc = Document(docx_file)
    return "\n".join([para.text for para in doc.paragraphs])


@st.cache_resource
def get_docx_text(docx_file):
    doc = Document(docx_file)
    return "\n".join([para.text for para in doc.paragraphs])


@st.cache_resource
def get_pdf_text(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


@st.cache_resource
def extract_images_from_pdf(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    images = []
    for page in pdf_reader.pages:
        if "/XObject" in page["/Resources"]:
            x_objects = page["/Resources"]["/XObject"].get_object()
            for obj in x_objects:
                if x_objects[obj]["/Subtype"] == "/Image":
                    try:
                        image_data = x_objects[obj].get_data()
                        image = PILImage.open(io.BytesIO(image_data))
                        images.append(image)
                    except Exception:
                        # Skip unsupported or corrupted images
                        continue
    return images


@st.cache_resource
def extract_images_from_docx(docx_file):
    doc = Document(docx_file)
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                image = PILImage.open(io.BytesIO(image_data))
                images.append(image)
            except Exception:
                # Skip unsupported or corrupted images
                continue
    return images